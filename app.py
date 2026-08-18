from flask import Flask, render_template, request, jsonify, send_file, Response, stream_with_context, redirect, url_for, flash
from flask_login import LoginManager, login_required, current_user, logout_user
from ai_brain import ask_leo_stream
from models import db, User, Note, QuizResult, RevisionSchedule
from auth import auth_bp
from sqlalchemy import or_
import json
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
import textwrap
import os
import uuid
import tempfile
import base64
import re
from html import unescape
import requests
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from datetime import datetime, timedelta

# ✅ File understanding — text/table extraction libs
import pdfplumber
from docx import Document as DocxDocument
import pandas as pd
from urllib.parse import urlparse, parse_qs

# ✅ YouTube transcript support (Notes Generator: YouTube URL -> Notes)
try:
    from youtube_transcript_api import YouTubeTranscriptApi
    from youtube_transcript_api._errors import (
        TranscriptsDisabled, NoTranscriptFound, VideoUnavailable
    )
except ImportError:
    # Package not installed yet — /fetch-youtube-transcript will return a
    # clear error instead of crashing the app at import time.
    YouTubeTranscriptApi = None
    TranscriptsDisabled = NoTranscriptFound = VideoUnavailable = Exception

# ✅ Web article extraction support (Notes Generator: Web article URL -> Notes)
try:
    from bs4 import BeautifulSoup
except ImportError:
    # Package not installed yet — /fetch-article-text falls back to a regex
    # text-stripper instead of crashing the app at import time.
    BeautifulSoup = None


app = Flask(__name__)

# =========================
# AUTH / DATABASE CONFIG
# =========================

app.config["SECRET_KEY"] = os.getenv("SECRET_KEY", "change-this-in-.env")
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///studyai.db"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

db.init_app(app)

login_manager = LoginManager()
login_manager.login_view = "auth.login"  # not-logged-in users get redirected here
login_manager.init_app(app)


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


app.register_blueprint(auth_bp)

with app.app_context():
    db.create_all()  # creates studyai.db + users table on first run


# =========================
# SPACED REVISION CONFIG
# =========================
# Days-until-next-review at each stage. A quiz score >=75% advances one
# stage (longer gap next time); 50-74% repeats the same stage; <50% resets
# to stage 0 (review again tomorrow). Caps out at the last stage rather than
# growing forever, so a "mastered" topic still resurfaces periodically.
REVISION_STAGE_INTERVALS_DAYS = [1, 3, 7, 14, 30]


# System instruction reused across chat turns — enables language auto-detect
# and keeps Leo's persona consistent across multi-turn conversation.
# NOTE: written in English (rather than Hinglish) and explicitly NOT limited
# to English/Hindi/Hinglish, so the model treats every language equally.
CHAT_SYSTEM_PROMPT = """You are "Leo", a friendly AI study assistant inside StudyAI, helping students
all over the world with their studies.

Important rule: Always reply in the SAME language and script the student used to ask their
question — this could be English, Hindi, Spanish, French, Arabic, Portuguese, Chinese, German,
Japanese, Russian, Bengali, Hinglish, or any other language. Detect it automatically, without
asking. If the user switches language mid-conversation, switch with them.

Keep answers clear, simple, and exam-oriented. Use bullet points, headings, or examples when
they help understanding.

If the user has attached a file, its extracted content will appear inline in their message
wrapped in an "[Attached file: ...]" block — treat it as ground truth context and answer based
on it. If web search results are provided in a system message, use them to give an up-to-date
answer and list the numbered sources at the end of your reply."""


# =========================
# FILE UPLOAD CONFIG
# =========================

ALLOWED_EXT = {
    "pdf", "docx", "txt", "csv",
    "png", "jpg", "jpeg",
    "py", "js", "java", "c", "cpp", "json", "md"
}
IMAGE_EXT = {"png", "jpg", "jpeg"}
MAX_UPLOAD_BYTES = 10 * 1024 * 1024  # 10 MB
MAX_EXTRACTED_CHARS = 15000  # keep token/cost usage under control


def extract_text_from_file(filepath, ext):
    """Extracts plain text from a non-image upload. Returns '' on anything
    unexpected rather than raising, so one bad file can't break the request."""
    try:
        if ext == "pdf":
            text = ""
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    text += (page.extract_text() or "") + "\n"
            return text
        elif ext == "docx":
            doc = DocxDocument(filepath)
            return "\n".join(p.text for p in doc.paragraphs)
        elif ext == "csv":
            df = pd.read_csv(filepath)
            return df.to_string(max_rows=200)  # truncate large CSVs
        elif ext in {"txt", "py", "js", "java", "c", "cpp", "json", "md"}:
            with open(filepath, "r", errors="ignore") as f:
                return f.read()
    except Exception:
        return ""
    return ""


# =========================
# WEB SEARCH CONFIG
# =========================

TAVILY_API_KEY = os.getenv("TAVILY_API_KEY", "").strip()
TAVILY_URL = "https://api.tavily.com/search"


def web_search(query, max_results=5):
    """Runs a Tavily search and returns a list of {title, content, url} dicts.
    Raises on failure — caller decides how to degrade gracefully."""
    if not TAVILY_API_KEY:
        raise ValueError("TAVILY_API_KEY not configured.")

    res = requests.post(
        TAVILY_URL,
        json={
            "api_key": TAVILY_API_KEY,
            "query": query,
            "max_results": max_results,
            "include_answer": False,
        },
        timeout=15,
    )
    res.raise_for_status()
    return res.json().get("results", [])


def build_web_context(results):
    """Formats search results into a compact numbered block for the model,
    plus a parallel list the frontend can render as clickable source chips."""
    lines = []
    sources = []
    for i, r in enumerate(results):
        title = r.get("title", "Untitled")
        url = r.get("url", "")
        content = (r.get("content") or "")[:500]
        lines.append(f"[{i + 1}] {title}\n{content}\nSource: {url}")
        sources.append({"title": title, "url": url})
    return "\n\n".join(lines), sources


@app.route("/")
def home():
    # current_user is available in templates automatically via Flask-Login,
    # no need to pass it explicitly — {{ current_user.is_authenticated }}
    # and {{ current_user.name }} just work inside index.html.
    return render_template("index.html")


@app.route("/privacy-policy")
def privacy_policy():
    return render_template("privacy.html")


@app.route("/sitemap.xml")
def sitemap():
    # NOTE: we only serve one real version of the page (no separate /?lang=hi
    # content actually exists server-side), so hreflang now only declares
    # "en" + "x-default" instead of falsely implying a dedicated Hindi page.
    # If/when true per-language routes are added, add matching hreflang
    # entries here pointing at those real URLs.
    sitemap_xml = """<?xml version="1.0" encoding="UTF-8"?>
<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9"
        xmlns:xhtml="http://www.w3.org/1999/xhtml">
  <url>
    <loc>https://zivolf.com/</loc>
    <changefreq>weekly</changefreq>
    <priority>1.0</priority>
    <xhtml:link rel="alternate" hreflang="en" href="https://zivolf.com/"/>
    <xhtml:link rel="alternate" hreflang="x-default" href="https://zivolf.com/"/>
  </url>
</urlset>"""
    return Response(sitemap_xml, mimetype="application/xml")


@app.route("/robots.txt")
def robots():
    robots_txt = """User-agent: *
Allow: /
Disallow: /ask-stream
Disallow: /generate-notes-stream
Disallow: /download-notes-pdf
Disallow: /generate-blog-stream
Disallow: /suggest-followups
Disallow: /generate-quiz
Disallow: /generate-flashcards
Disallow: /generate-mindmap
Disallow: /upload-file
Disallow: /fetch-youtube-transcript
Disallow: /fetch-article-text
Disallow: /notes-tutor-stream
Disallow: /pdf-workspace-stream
Disallow: /api/notes
Disallow: /api/notes/search
Disallow: /api/export-data
Disallow: /account/delete
Disallow: /api/quiz-results
Disallow: /dashboard
Disallow: /download-notes-docx
Disallow: /shared-notes/
Disallow: /static/generated/

Sitemap: https://zivolf.com/sitemap.xml"""
    return Response(robots_txt, mimetype="text/plain")


@app.route("/ads.txt")
def ads_txt():
    return app.send_static_file("ads.txt")


# =========================
# FILE UPLOAD (used by AI Chat attachments AND Notes Generator source files)
# =========================

@app.route("/upload-file", methods=["POST"])
def upload_file():
    if "file" not in request.files:
        return jsonify({"error": "File nahi mili."}), 400

    file = request.files["file"]
    if not file.filename:
        return jsonify({"error": "File nahi mili."}), 400

    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ALLOWED_EXT:
        return jsonify({"error": f"Unsupported file type: .{ext}"}), 400

    # Size check before saving to disk
    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)
    if size > MAX_UPLOAD_BYTES:
        return jsonify({"error": "File 10MB se badi hai."}), 400

    tmp_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex}_{file.filename}")
    file.save(tmp_path)

    try:
        if ext in IMAGE_EXT:
            with open(tmp_path, "rb") as f:
                b64 = base64.b64encode(f.read()).decode()
            mime = "jpeg" if ext == "jpg" else ext
            return jsonify({
                "type": "image",
                "filename": file.filename,
                "data_url": f"data:image/{mime};base64,{b64}"
            })
        else:
            text = extract_text_from_file(tmp_path, ext)
            if not text.strip():
                return jsonify({"error": "File se text extract nahi ho paya."}), 400
            return jsonify({
                "type": "text",
                "filename": file.filename,
                "content": text[:MAX_EXTRACTED_CHARS]
            })
    finally:
        try:
            os.remove(tmp_path)
        except OSError:
            pass


# =========================
# YOUTUBE TRANSCRIPT (used by Notes Generator's "YouTube URL" source)
# =========================

MAX_TRANSCRIPT_CHARS = 15000  # mirrors MAX_EXTRACTED_CHARS for uploaded files


def extract_youtube_id(url):
    """Pulls the 11-char video ID out of common YouTube URL shapes:
    watch?v=, youtu.be/, /embed/, /shorts/, /live/. Returns None if the
    URL doesn't look like YouTube at all."""
    url = (url or "").strip()
    if not url:
        return None

    parsed = urlparse(url)
    host = (parsed.hostname or "").lower()

    if "youtu.be" in host:
        vid = parsed.path.lstrip("/").split("/")[0]
        return vid or None

    if "youtube.com" in host:
        if parsed.path == "/watch":
            qs = parse_qs(parsed.query)
            return (qs.get("v", [None])[0]) or None
        for prefix in ("/embed/", "/shorts/", "/live/"):
            if parsed.path.startswith(prefix):
                vid = parsed.path[len(prefix):].split("/")[0]
                return vid or None

    return None


@app.route("/fetch-youtube-transcript", methods=["POST"])
def fetch_youtube_transcript():
    data = request.get_json(silent=True) or {}
    url = data.get("url", "").strip()

    if not url:
        return jsonify({"error": "YouTube URL khaali hai."}), 400

    if YouTubeTranscriptApi is None:
        return jsonify({
            "error": "YouTube support abhi server par install nahi hai "
                     "(pip install youtube-transcript-api)."
        }), 500

    video_id = extract_youtube_id(url)
    if not video_id:
        return jsonify({"error": "Yeh ek valid YouTube video URL nahi lag raha."}), 400

    try:
        transcript_list = YouTubeTranscriptApi.get_transcript(video_id)
        text = " ".join(seg.get("text", "") for seg in transcript_list).strip()
        if not text:
            raise ValueError("Empty transcript")
    except TranscriptsDisabled:
        return jsonify({"error": "Is video par captions/transcript disabled hain."}), 400
    except NoTranscriptFound:
        return jsonify({"error": "Is video ka transcript nahi mila."}), 400
    except VideoUnavailable:
        return jsonify({"error": "Video available nahi hai ya private/restricted hai."}), 400
    except Exception:
        return jsonify({"error": "Transcript fetch nahi ho paya, dobara try karein."}), 400

    # Same {type, filename, content} shape /upload-file returns for text
    # files, plus a "source" hint the frontend uses to pick a 🎥 icon —
    # this lets it reuse the exact same single-slot attachment mechanism.
    return jsonify({
        "type": "text",
        "source": "youtube",
        "filename": f"YouTube video ({video_id})",
        "content": text[:MAX_TRANSCRIPT_CHARS]
    })


# =========================
# WEB ARTICLE TEXT (used by Notes Generator's "Web article URL" source)
# =========================

ARTICLE_FETCH_TIMEOUT = 15  # seconds — mirrors web_search()'s timeout
ARTICLE_USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
    "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"
)
# Boilerplate containers never carry article content — strip these subtrees
# before extracting text so notes don't get polluted with nav/footer noise.
ARTICLE_BOILERPLATE_TAGS = (
    "script", "style", "noscript", "nav", "footer", "header",
    "aside", "form", "iframe", "svg", "canvas", "template", "dialog",
)


def _extract_article_content(raw_html):
    """Returns (text, title) for a fetched page's HTML.

    Prefers BeautifulSoup (beautifulsoup4 in requirements.txt) — it can
    properly decompose boilerplate subtrees and prefers the semantic
    <article> region when a page marks one. When the package isn't installed
    yet, falls back to a regex-based stripper: cruder, but it still drops
    the same noise blocks and pulls readable text out."""
    if not raw_html:
        return "", None

    if BeautifulSoup is not None:
        soup = BeautifulSoup(raw_html, "html.parser")
        for name in ARTICLE_BOILERPLATE_TAGS:
            for node in soup.find_all(name):
                node.decompose()
        # Prefer the article block; otherwise the whole body.
        root = soup.find("article") or soup.body or soup
        text = "\n".join(
            line.strip() for line in root.get_text("\n", strip=True).splitlines()
            if line.strip()
        )
        title = None
        if soup.title and soup.title.get_text(strip=True):
            title = " ".join(soup.title.get_text(" ", strip=True).split())
        return text, title

    # Regex fallback (beautifulsoup4 not installed yet).
    stripped = raw_html
    for tag in ARTICLE_BOILERPLATE_TAGS:
        stripped = re.sub(
            rf"<{tag}\b[^>]*>.*?</{tag}>", " ", stripped,
            flags=re.IGNORECASE | re.DOTALL,
        )
    stripped = re.sub(r"<!--.*?-->", " ", stripped, flags=re.DOTALL)
    stripped = re.sub(r"<[^>]+>", " ", stripped)
    stripped = unescape(stripped)  # &amp; &lt; ... back to readable characters
    text = "\n".join(
        line.strip() for line in stripped.splitlines() if line.strip()
    )

    title = None
    m = re.search(r"<title\b[^>]*>(.*?)</title>", raw_html,
                  flags=re.IGNORECASE | re.DOTALL)
    if m:
        title = re.sub(r"<[^>]+>", " ", m.group(1)).strip()
        title = unescape(" ".join(title.split())) or None
    return text, title


def _article_filename(title, url):
    """Uses the page <title> when we got one, else the URL — kept short
    enough to fit neatly in the attachment chip."""
    label = (title or url or "Web article").strip()
    if not label:
        label = "Web article"
    return label if len(label) <= 120 else label[:117].rstrip() + "..."


@app.route("/fetch-article-text", methods=["POST"])
def fetch_article_text():
    data = request.get_json(silent=True) or {}
    url = data.get("url", "").strip()

    if not url:
        return jsonify({"error": "Article URL khaali hai."}), 400

    # Only http/https — reject javascript:, data:, file: etc. outright.
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https") or not parsed.netloc:
        return jsonify({
            "error": "Yeh ek valid web page URL nahi lag raha (http/https chahiye)."
        }), 400

    try:
        resp = requests.get(
            url,
            headers={"User-Agent": ARTICLE_USER_AGENT},
            timeout=ARTICLE_FETCH_TIMEOUT,
        )
    except requests.exceptions.Timeout:
        return jsonify({
            "error": "URL ka response aane me time lag gaya (timeout) — dobara try karein."
        }), 400
    except requests.exceptions.RequestException:
        return jsonify({
            "error": "Is URL tak pahunch nahi paye — URL galat, unreachable, ya site ne request block kar di."
        }), 400

    if resp.status_code != 200:
        reason = (
            "site ne scraping block kar diya hai" if resp.status_code in (401, 403)
            else "page exist nahi karta" if resp.status_code == 404
            else "server ne error return kiya"
        )
        return jsonify({
            "error": f"Page ne status {resp.status_code} return kiya ({reason}) — content extract nahi ho paya."
        }), 400

    # Some sites serve content with a mismatched charset header — sniff it.
    resp.encoding = resp.apparent_encoding or resp.encoding

    text, title = _extract_article_content(resp.text)
    if not text:
        return jsonify({
            "error": "Is URL se content nahi mila — page par extractable text nahi hai."
        }), 400

    # Same {type, filename, content} shape file/YouTube/paste sources use,
    # plus a "source" hint so the frontend chip can show a 🌐 icon.
    return jsonify({
        "type": "text",
        "source": "article",
        "filename": _article_filename(title, url),
        "content": text[:MAX_EXTRACTED_CHARS]
    })


# =========================
# AI CHAT (STREAMING, MULTI-TURN)
# =========================

@app.route("/ask-stream", methods=["POST"])
# @login_required   # <-- uncomment this if you want chat to require login.
                     # Left open for now so guests can keep using chat
                     # without an account, matching the current homepage copy.
def ask_stream():
    data = request.get_json(silent=True) or {}
    question = data.get("question", "").strip()
    history = data.get("history", [])  # [{role, content}, ...] sent by frontend, includes current turn
    attachments = data.get("attachments", [])  # [{type, filename, content|data_url}, ...]
    web_search_enabled = bool(data.get("web_search", False))

    if not question:
        return jsonify({"error": "Sawal khaali hai."}), 400

    if not history:
        history = [{"role": "user", "content": question}]

    # Cap history length sent to the model to control token/cost growth
    MAX_TURNS = 20
    trimmed_history = history[-MAX_TURNS:]

    # --- Fold attachments into the current (last) user turn ---
    # Text-file attachments get appended as inline context; image attachments
    # turn that turn's content into a multimodal block (text + image_url parts).
    image_attachments = [a for a in attachments if a.get("type") == "image"]
    text_attachments = [a for a in attachments if a.get("type") == "text"]

    if trimmed_history and trimmed_history[-1].get("role") == "user":
        last_turn = dict(trimmed_history[-1])
        base_text = last_turn.get("content", question)

        for att in text_attachments:
            base_text += f"\n\n[Attached file: {att.get('filename', 'file')}]\n{att.get('content', '')}"

        if image_attachments:
            content_blocks = [{"type": "text", "text": base_text}]
            for att in image_attachments:
                content_blocks.append({
                    "type": "image_url",
                    "image_url": {"url": att.get("data_url", "")}
                })
            last_turn["content"] = content_blocks
        else:
            last_turn["content"] = base_text

        trimmed_history = trimmed_history[:-1] + [last_turn]

    messages = [{"role": "system", "content": CHAT_SYSTEM_PROMPT}] + trimmed_history

    web_sources = []
    if web_search_enabled:
        try:
            results = web_search(question)
            context, web_sources = build_web_context(results)
            if context:
                messages.append({
                    "role": "system",
                    "content": (
                        "Here are fresh web search results — use them to answer with "
                        "up-to-date information, and list the numbered sources at the "
                        "end of your reply:\n\n" + context
                    )
                })
        except Exception:
            # Search failing shouldn't block the chat — fall back to a normal answer.
            pass

    def generate():
        try:
            # Send sources first so the frontend can render them even before
            # the model finishes streaming its answer.
            if web_sources:
                yield f"data: {json.dumps({'sources': web_sources})}\n\n"
            for chunk in ask_leo_stream(messages=messages, stream=True):
                yield f"data: {json.dumps({'chunk': chunk})}\n\n"
            yield f"data: {json.dumps({'done': True})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return Response(stream_with_context(generate()), mimetype="text/event-stream")


# =========================
# FOLLOW-UP QUESTION SUGGESTIONS
# =========================

@app.route("/suggest-followups", methods=["POST"])
def suggest_followups():
    data = request.get_json(silent=True) or {}
    question = data.get("question", "").strip()
    answer = data.get("answer", "").strip()

    if not question or not answer:
        return jsonify({"suggestions": []})

    prompt = f"""Below is a Q&A pair. Suggest exactly 3 short, natural follow-up questions a
student might ask next — write them in the SAME language/script as the question below
(whatever language that is).

Reply ONLY as a raw JSON array of exactly 3 short strings. No extra text, no markdown, no code fences.

Question: {question}
Answer: {answer[:800]}
"""

    try:
        parts = list(ask_leo_stream(
            messages=[{"role": "user", "content": prompt}],
            stream=False
        ))
        raw = "".join(parts).strip()

        # Defensive cleanup in case the model wraps the JSON in code fences
        if raw.startswith("```"):
            raw = raw.strip("`")
            if raw.lower().startswith("json"):
                raw = raw[4:].strip()

        suggestions = json.loads(raw)
        if not isinstance(suggestions, list):
            suggestions = []
    except Exception:
        suggestions = []

    return jsonify({"suggestions": suggestions[:3]})


# =========================
# SHARED JSON-RESPONSE HELPERS
# (used by quiz + flashcards + mind map, all of which ask the model to
#  reply as raw JSON — models sometimes wrap that in code fences anyway)
# =========================

def _parse_json_response(raw):
    """Parses a raw model response as JSON, tolerating a markdown
    code-fence wrapper (models sometimes wrap JSON in ```json ... ```)."""
    raw = raw.strip()
    if raw.startswith("```"):
        raw = raw.strip("`")
        if raw.lower().startswith("json"):
            raw = raw[4:].strip()
    return json.loads(raw)


# =========================
# AI NOTES GENERATOR (STREAMING)
# =========================

# Available note shapes toggled by the frontend's "Note Style" dropdown.
# Each entry maps a style key to (goal_line, format_block): the opening
# intent shown above the topic fields, and the "Notes format" section that
# shapes the output. "detailed" is the long-standing default — its block must
# stay byte-identical so existing behaviour never changes unless the user
# deliberately picks another style.
NOTE_STYLES = {
    "detailed": (
        "Create the best possible exam-ready study notes for a student.",
        (
            "Notes format:\n"
            "1. Short Introduction\n"
            "2. Simple Explanation\n"
            "3. Important Definitions\n"
            "4. Important Points\n"
            "5. Examples\n"
            "6. Exam Tips\n"
            "7. 10 MCQ Questions with Answers\n"
            "8. Short Summary\n"
            "\n"
            "Keep the notes simple, clear, and exam-ready.\n"
        ),
    ),
    "quick": (
        "Create quick, bite-sized study notes that highlight the core facts.",
        (
            "Notes format:\n"
            "- Short bullet points only\n"
            "- Key facts and essential takeaways\n"
            "- No long explanations or paragraphs\n"
            "\n"
            "Keep every bullet short, direct, and easy to skim.\n"
        ),
    ),
    "exam": (
        "Create exam-focused study notes that get straight to what is likely to be asked.",
        (
            "Notes format:\n"
            "1. Important Questions\n"
            "2. Key Concepts\n"
            "3. Likely Exam Questions (with short answers)\n"
            "4. Formulas / Rules / Diagrams to remember\n"
            "\n"
            "Write everything from an examiner's point of view — precise, high-yield, no fluff.\n"
        ),
    ),
    "revision": (
        "Create a tight last-minute revision summary covering only what must be remembered.",
        (
            "Notes format:\n"
            "1. Core idea in a line or two\n"
            "2. Must-know points (compact)\n"
            "3. Key formulas / terms / dates\n"
            "4. Tricky spots or common mistakes\n"
            "\n"
            "Keep it to a tight 1-2 page summary that can be re-read in minutes before an exam.\n"
        ),
    ),
    "beginner": (
        "Create study notes for a complete beginner with zero background in this topic.",
        (
            "Notes format:\n"
            "1. Start from the very basics\n"
            "2. Explain every idea in simple, everyday words\n"
            "3. Avoid jargon — when a technical word is needed, define it simply\n"
            "4. Use short sentences and friendly examples\n"
            "\n"
            "Explain like a patient teacher starting from scratch; nothing should feel assumed.\n"
        ),
    ),
    "professional": (
        "Create technical, precise study notes for working professionals or researchers.",
        (
            "Notes format:\n"
            "1. Technical Overview\n"
            "2. Key Concepts (exact terminology)\n"
            "3. Detailed Explanations of Core Mechanisms\n"
            "4. Formulas / Models / Specifications\n"
            "5. Caveats, Edge Cases, and Practical Notes\n"
            "\n"
            "Use precise, domain-accurate terminology. Prioritize depth and correctness over simplicity.\n"
        ),
    ),
}


# =========================
# NOTES GENERATOR MODES (student | teacher | professional)
# The frontend sends one `mode` field; /generate-notes-stream swaps the
# student "Notes format:" block for a lesson-plan package (teacher) or a
# workplace brief (professional). All three modes reuse the exact same
# generation flow / streaming / PDF export / history-save code paths — only
# the prompt format changes.
# =========================

TEACHER_MODE_GOAL = (
    "Create a complete, ready-to-use lesson-plan package for a teacher who will teach this "
    "topic to a class."
)

TEACHER_MODE_FORMAT = (
    "Lesson-package format:\n"
    "1. Lesson objectives (what students should be able to do by the end of the lesson)\n"
    "2. Teaching outline / lesson plan with rough timing (e.g. 5 min warm-up, 15 min explanation, "
    "10 min activity, 5 min wrap-up)\n"
    "3. Key points to emphasize while teaching\n"
    "4. Suggested in-class activities / discussion questions\n"
    "5. Homework assignment\n"
    "6. A short worksheet (a handful of practice questions)\n"
    "7. An answer key for the worksheet (with brief explanations)\n"
    "\n"
    "Make it practical and ready to use in a real classroom. The lesson plan, key points, "
    "activities, and answer key are for the teacher's eyes; the homework and worksheet are "
    "given to students.\n"
)

PROFESSIONAL_MODE_GOAL = (
    "Treat the provided input as workplace content (a meeting transcript, a research article, "
    "a report, or general work notes) and produce a concise, action-oriented workplace brief."
)

PROFESSIONAL_MODE_FORMAT = (
    "Workplace-brief format:\n"
    "1. Executive summary (one concise paragraph capturing the essence of the material)\n"
    "2. Key points / decisions made (a short bullet list)\n"
    "3. Action items (as a checklist — include the owner's name for each item when the source "
    "text names one)\n"
    "4. Risks or open questions worth flagging\n"
    "5. Suggested next steps\n"
    "\n"
    "Write in clear, professional business language — concise, precise, and actionable. "
    "This is a workplace brief, not student study notes.\n"
)


def _build_notes_prompt(topic, subject, class_name, language, has_attachment,
                        note_style="detailed", mode="student"):
    """Builds the base instruction text for notes generation. Kept separate
    from message/content assembly so both the plain-text path and the
    image (multimodal) path can share identical wording. `note_style`
    selects which student "Notes format" block shapes the output; unknown
    styles are normalized to "detailed" upstream so this builder can trust
    its input.

    `mode` picks the output audience/shape:
      - "student" (default): exam-ready study notes (Note Style applies).
        Byte-identical to the long-standing behaviour.
      - "teacher": a ready-to-use lesson-plan package (TEACHER_MODE_FORMAT).
      - "professional": a workplace brief (PROFESSIONAL_MODE_FORMAT).
    The latter two are fixed formats, so `note_style` is ignored there.
    Unknown modes are normalized to "student" upstream."""
    topic_label = topic if topic else "the attached material"

    if mode == "teacher":
        # Teacher Mode — lesson-plan package instead of student study notes.
        source_note = (
            "\nSource material (one or more documents, images, transcripts, or web "
            "articles) is attached — base the lesson-plan package primarily on that "
            "attached material. If several sources are attached, COMBINE all of them "
            "into one coherent package. Use the topic/subject/class fields below only "
            "as extra guidance if they were provided; if the topic wasn't typed, "
            "infer a sensible topic/title from the attached content instead.\n"
            if has_attachment else ""
        )
        goal_line = TEACHER_MODE_GOAL
        format_block = TEACHER_MODE_FORMAT
    elif mode == "professional":
        # Professional Mode — workplace brief instead of student study notes.
        source_note = (
            "\nSource material (one or more documents, images, transcripts, or web "
            "articles) is attached — base the workplace brief primarily on that "
            "attached material. If several sources are attached, COMBINE all of them "
            "into one coherent brief. Use the topic/subject/context fields below only "
            "as extra guidance if they were provided; if the topic wasn't typed, "
            "infer a sensible topic/title from the attached content instead.\n"
            if has_attachment else ""
        )
        goal_line = PROFESSIONAL_MODE_GOAL
        format_block = PROFESSIONAL_MODE_FORMAT
    else:
        # Student Mode — the long-standing default, byte-identical.
        source_note = (
            "\nSource material (one or more documents, images, transcripts, or web "
            "articles) is attached — base the notes primarily on that attached "
            "material. If several sources are attached, COMBINE all of them into one "
            "coherent, unified set of notes. Use the topic/subject/class fields below "
            "only as extra guidance if they were provided; if the topic wasn't typed, "
            "infer a sensible topic/title from the attached content instead.\n"
            if has_attachment else ""
        )

        # The goal line drives the overall voice, while the format block fixes
        # the exact structure. Fall back to "detailed" just in case a style key
        # slips through validation.
        goal_line, format_block = NOTE_STYLES.get(note_style, NOTE_STYLES["detailed"])

    if mode == "professional":
        persona_line = (
            "You are an experienced professional who turns meetings, reports, and research "
            "into concise, action-oriented workplace briefs."
        )
    else:
        persona_line = "You are an expert teacher."

    return f"""
{persona_line}

{source_note}{goal_line}
Topic: {topic_label}
Subject: {subject}
Class/Level: {class_name}
Language: {language}

Write the ENTIRE output in the language specified above — this could be English, Hindi,
Spanish, French, Arabic, Portuguese, Chinese, German, Japanese, Russian, Bengali, or any other
language. Follow it exactly, including using that language's own script.

{format_block}"""

# Combined character budget for ALL text attachments in one notes request.
# Every individual source is already capped at MAX_EXTRACTED_CHARS (15,000);
# with up to 5 sources that could reach ~75k chars (~20k+ tokens), far too big
# for a single prompt. 2x MAX_EXTRACTED_CHARS keeps a multi-document run
# economical (~7-8k tokens) while still comfortably fitting 2-3 full extracted
# chapters of study material.
MAX_NOTES_TEXT_ATTACHMENT_CHARS = MAX_EXTRACTED_CHARS * 2  # 30,000


def _is_usable_attachment(att):
    """True for attachment dicts that actually carry something usable (text
    content or an image data URL), so junk/malformed entries can't leak into
    the prompt."""
    if not isinstance(att, dict):
        return False
    if att.get("type") == "text":
        return bool(att.get("content"))
    if att.get("type") == "image":
        return bool(att.get("data_url"))
    return False


def _format_text_attachments(text_attachments):
    """Labels and concatenates every text source under one shared character
    budget, e.g.
        [Attached source 1: chapter1.pdf]\n...
        \n[Attached source 2: chapter2.pdf]\n...
    Each source keeps its own label even when the budget forces truncation."""
    used, blocks = 0, []
    for i, att in enumerate(text_attachments, start=1):
        content = att.get("content", "") or ""
        if not content:
            continue
        label = att.get("filename") or f"source {i}"
        remaining = MAX_NOTES_TEXT_ATTACHMENT_CHARS - used
        if remaining <= 0:
            break
        if len(content) > remaining:
            blocks.append(f"[Attached source {i}: {label}]\n{content[:remaining]}")
            used = MAX_NOTES_TEXT_ATTACHMENT_CHARS
            break
        blocks.append(f"[Attached source {i}: {label}]\n{content}")
        used += len(content)
    return "\n\n".join(blocks)


@app.route("/generate-notes-stream", methods=["POST"])
def generate_notes_stream():
    data = request.get_json(silent=True) or {}
    topic = data.get("topic", "").strip()
    subject = data.get("subject", "").strip()
    class_name = data.get("class_name", "").strip()
    # Global default is English rather than Hindi, since this app now
    # targets students worldwide, not just Hindi-speaking students.
    language = data.get("language", "English").strip()
    # Optional source material staged by the Notes Generator: an array of
    # {type:'text'|'image', source?, filename, content|data_url} entries from
    # files, images, YouTube transcripts, pasted text, and web articles. The
    # old singular "attachment" key is still accepted (wrapped in a 1-item
    # list) so a cached older frontend mid-transition keeps working.
    raw_attachments = data.get("attachments")
    if not isinstance(raw_attachments, list):
        raw_attachments = [data.get("attachment")] if data.get("attachment") else []
    attachments = [a for a in raw_attachments if _is_usable_attachment(a)]
    # "Note Style" shape picked in the frontend (detailed/quick/exam/
    # revision/beginner/professional). Unknown or missing values fall back
    # to "detailed" — old cached frontends never send this field, so they
    # must keep working without erroring.
    raw_style = data.get("note_style") or "detailed"
    note_style = raw_style.strip() if isinstance(raw_style, str) else "detailed"
    if note_style not in NOTE_STYLES:
        note_style = "detailed"
    # Notes mode: one clean field selects the output audience/shape.
    #   "student"      (default) — exam-ready study notes (Note Style applies)
    #   "teacher"                — lesson-plan package
    #   "professional"           — workplace brief
    # Unknown/missing values fall back to "student", so old cached frontends
    # that predate the field keep working without erroring.
    raw_mode = data.get("mode") or "student"
    mode = raw_mode.strip() if isinstance(raw_mode, str) else "student"
    if mode not in ("student", "teacher", "professional"):
        mode = "student"

    if not topic and not attachments:
        return jsonify({"error": "Topic ya attachment me se ek dena zaroori hai."}), 400

    prompt = _build_notes_prompt(
        topic, subject, class_name, language,
        has_attachment=bool(attachments), note_style=note_style,
        mode=mode,
    )

    text_attachments = [a for a in attachments if a.get("type") == "text"]
    image_attachments = [a for a in attachments if a.get("type") == "image"]

    # Fold every text source into the prompt, each labelled, under one shared
    # character budget (see MAX_NOTES_TEXT_ATTACHMENT_CHARS).
    if text_attachments:
        prompt += "\n\n" + _format_text_attachments(text_attachments)

    if image_attachments:
        # Multimodal — the same shape /ask-stream uses for chat: all images as
        # separate image_url blocks inside one user message, so a vision model
        # sees every image (e.g. photos of several textbook pages at once).
        content_blocks = [{"type": "text", "text": prompt}]
        for att in image_attachments:
            content_blocks.append({
                "type": "image_url",
                "image_url": {"url": att.get("data_url", "")}
            })
        messages = [{"role": "user", "content": content_blocks}]
    else:
        messages = [{"role": "user", "content": prompt}]

    def generate():
        try:
            for chunk in ask_leo_stream(messages=messages, stream=True):
                yield f"data: {json.dumps({'chunk': chunk})}\n\n"
            yield f"data: {json.dumps({'done': True})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return Response(stream_with_context(generate()), mimetype="text/event-stream")


# =========================
# AI QUIZ GENERATOR (from Notes)
# =========================

@app.route("/generate-quiz", methods=["POST"])
def generate_quiz():
    data = request.get_json(silent=True) or {}
    notes = data.get("notes", "").strip()
    language = data.get("language", "English").strip()

    if not notes:
        return jsonify({"error": "Notes khaali hai."}), 400

    prompt = f"""You are an expert teacher. Based on the following study notes, create exactly 8 quiz
questions that test understanding of the material. Use a MIX of question types in one quiz:
at least 4 multiple-choice ("mcq"), 1-2 true/false ("true_false"), and 1-2 fill-in-the-blank
("fill_blank") questions.

Notes:
{notes[:6000]}

Reply ONLY as a raw JSON array (no markdown, no code fences, no extra text) of exactly 8 objects.
Every object must include a "type" field, one of: "mcq" | "true_false" | "fill_blank".

For "mcq" questions, include exactly these keys:
- "question": string
- "options": array of exactly 4 short strings
- "answer": integer index (0-3) of the correct option in "options"
- "explanation": one short sentence explaining why that answer is correct

For "true_false" questions, include exactly these keys:
- "question": string that can be answered True or False
- "options": exactly ["True", "False"]
- "answer": 0 for True or 1 for False
- "explanation": one short sentence explaining why that answer is correct

For "fill_blank" questions, include exactly these keys:
- "question": string containing a blank written as ____
- "answer_text": the correct short answer as a string (no "options" or "answer" keys)
- "explanation": one short sentence explaining the correct answer

Write all text in this language: {language}.
"""

    try:
        parts = list(ask_leo_stream(
            messages=[{"role": "user", "content": prompt}],
            stream=False
        ))
        raw = "".join(parts)
        questions = _parse_json_response(raw)
        if not isinstance(questions, list) or not questions:
            raise ValueError("Empty quiz response")

        # Basic shape validation — drop malformed items (per type) rather than
        # failing the whole quiz, mirroring the old MCQ-only defensive style.
        clean = []
        for q in questions:
            if not isinstance(q, dict) or not isinstance(q.get("question"), str) or not q["question"].strip():
                continue

            # Older responses may omit "type" — default missing/unknown types
            # to multiple-choice so nothing from before this feature breaks.
            qtype = q.get("type")
            if qtype not in ("mcq", "true_false", "fill_blank"):
                qtype = "mcq"

            if qtype == "fill_blank":
                answer_text = q.get("answer_text")
                if not isinstance(answer_text, str) or not answer_text.strip():
                    continue  # a blank without a known answer is useless — drop
                clean.append({
                    "type": "fill_blank",
                    "question": q["question"],
                    "answer_text": answer_text.strip(),
                    "explanation": str(q.get("explanation", ""))
                })
                continue

            # mcq + true_false share the options/answer-as-index shape.
            options = q.get("options")
            answer = q.get("answer")
            if not isinstance(options, list) or len(options) < 2:
                continue
            if isinstance(answer, bool):  # JSON true/false pops out as a Python bool
                answer = int(answer)
            if not isinstance(answer, int) or not (0 <= answer < len(options)):
                continue

            if qtype == "mcq":
                clean.append({
                    "type": "mcq",
                    "question": q["question"],
                    "options": [str(o) for o in options[:6]],
                    "answer": answer,
                    "explanation": str(q.get("explanation", ""))
                })
            else:  # true_false
                if answer not in (0, 1):
                    continue  # true/false answer must be exactly 0 or 1
                clean.append({
                    "type": "true_false",
                    "question": q["question"],
                    "options": ["True", "False"],  # always normalize the labels
                    "answer": answer,
                    "explanation": str(q.get("explanation", ""))
                })

        if not clean:
            raise ValueError("No valid quiz questions")

        return jsonify({"questions": clean})
    except Exception:
        return jsonify({"error": "Quiz nahi ban paya, dobara try karein."}), 500


# =========================
# AI FLASHCARDS GENERATOR (from Notes)
# =========================

@app.route("/generate-flashcards", methods=["POST"])
def generate_flashcards():
    data = request.get_json(silent=True) or {}
    notes = data.get("notes", "").strip()
    language = data.get("language", "English").strip()

    if not notes:
        return jsonify({"error": "Notes khaali hai."}), 400

    prompt = f"""You are an expert teacher. Based on the following study notes, create exactly 12
flashcards that help a student memorize and revise the key facts, terms, and concepts.

Notes:
{notes[:6000]}

Reply ONLY as a raw JSON array (no markdown, no code fences, no extra text) of exactly 12 objects,
each with these exact keys:
- "front": a short question, term, or prompt (max ~12 words)
- "back": the concise answer or definition (max ~30 words)

Write all text in this language: {language}.
"""

    try:
        parts = list(ask_leo_stream(
            messages=[{"role": "user", "content": prompt}],
            stream=False
        ))
        raw = "".join(parts)
        cards = _parse_json_response(raw)
        if not isinstance(cards, list) or not cards:
            raise ValueError("Empty flashcards response")

        clean = [
            {"front": str(c["front"]), "back": str(c["back"])}
            for c in cards
            if isinstance(c, dict) and isinstance(c.get("front"), str) and isinstance(c.get("back"), str)
        ]

        if not clean:
            raise ValueError("No valid flashcards")

        return jsonify({"cards": clean})
    except Exception:
        return jsonify({"error": "Flashcards nahi ban paye, dobara try karein."}), 500


# =========================
# AI MIND MAP GENERATOR (from Notes)
# Returns a 2-level tree {topic, branches:[{label, children:[...]}]} that
# the frontend renders as an SVG radial diagram — see renderMindMap() in
# index.html. Not persisted to the database; generated fresh each time,
# same as Quiz and Flashcards.
# =========================

@app.route("/generate-mindmap", methods=["POST"])
def generate_mindmap():
    data = request.get_json(silent=True) or {}
    notes = data.get("notes", "").strip()
    language = data.get("language", "English").strip()

    if not notes:
        return jsonify({"error": "Notes khaali hai."}), 400

    prompt = f"""You are an expert teacher. Based on the following study notes, design a mind map
that captures the topic's structure at a glance.

Notes:
{notes[:6000]}

Reply ONLY as a raw JSON object (no markdown, no code fences, no extra text) with this exact shape:
{{
  "topic": "central topic, a few words",
  "branches": [
    {{
      "label": "main branch label, a few words",
      "children": ["short sub-point", "short sub-point"]
    }}
  ]
}}

Rules:
- 4 to 7 branches total, covering the main themes of the notes.
- Each branch has 2 to 5 short children (a few words each, not full sentences).
- Keep every label and child short and scannable — this is a visual diagram, not prose.
- Write all text in this language: {language}.
"""

    try:
        parts = list(ask_leo_stream(
            messages=[{"role": "user", "content": prompt}],
            stream=False
        ))
        raw = "".join(parts)
        parsed = _parse_json_response(raw)

        if not isinstance(parsed, dict):
            raise ValueError("Not a JSON object")

        topic = str(parsed.get("topic") or "Mind Map").strip()[:80]
        raw_branches = parsed.get("branches")
        if not isinstance(raw_branches, list) or not raw_branches:
            raise ValueError("No branches")

        clean_branches = []
        for b in raw_branches[:8]:
            if not isinstance(b, dict):
                continue
            label = str(b.get("label") or "").strip()[:60]
            if not label:
                continue

            children = []
            children_raw = b.get("children")
            if isinstance(children_raw, list):
                for c in children_raw[:6]:
                    c_str = str(c).strip()[:60]
                    if c_str:
                        children.append(c_str)

            clean_branches.append({"label": label, "children": children})

        if not clean_branches:
            raise ValueError("No valid branches")

        return jsonify({"topic": topic, "branches": clean_branches})
    except Exception:
        return jsonify({"error": "Mind map nahi ban paya, dobara try karein."}), 500


# =========================
# AI TUTOR — "Ask AI About These Notes" (STREAMING)
# A lightweight, separate Q&A grounded in the current Notes box. Does NOT
# touch CHAT_SYSTEM_PROMPT or the main AI Chat's multi-turn history — this
# is its own scoped, single-turn-per-question conversation (frontend keeps
# any back-and-forth purely as local chat bubbles, not as server state).
# =========================

@app.route("/notes-tutor-stream", methods=["POST"])
def notes_tutor_stream():
    data = request.get_json(silent=True) or {}
    question = data.get("question", "").strip()
    notes = data.get("notes", "").strip()
    language = data.get("language", "English").strip()

    if not question:
        return jsonify({"error": "Sawal khaali hai."}), 400
    if not notes:
        return jsonify({"error": "Pehle notes generate karein."}), 400

    prompt = f"""You are a friendly, patient AI tutor helping a student understand their own study
notes below. Stay grounded in these notes — you may add a little extra context or a helpful
analogy, but don't contradict what the notes say.

Notes:
{notes[:6000]}

The student's question about these notes:
{question}

Answer clearly and simply. Use examples, analogies, or a short step-by-step breakdown where it
helps understanding. Reply in this language: {language}.
"""

    def generate():
        try:
            for chunk in ask_leo_stream(messages=[{"role": "user", "content": prompt}], stream=True):
                yield f"data: {json.dumps({'chunk': chunk})}\n\n"
            yield f"data: {json.dumps({'done': True})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return Response(stream_with_context(generate()), mimetype="text/event-stream")


# =========================
# PDF WORKSPACE (already-extracted PDF text + an instruction)
# Summarize / Ask / Explain / Translate / Compare — a dedicated
# workflow around capability that already exists (attach a file's
# extracted text as context for a model call). Reuses /upload-file to
# get the text into the frontend; this route only builds a prompt and
# streams the answer, mirroring /notes-tutor-stream and /ask-stream.
# =========================

def _build_pdf_workspace_prompt(action, pdf_text_1, filename_1,
                                pdf_text_2=None, filename_2=None,
                                question="", target_language=""):
    """Builds the instruction text for the PDF Workspace, branching per
    action the same way _build_notes_prompt branches by note style, so the
    route itself never assembles prompt strings inline. Each document is
    already capped at MAX_EXTRACTED_CHARS upstream (/upload-file); we slice
    defensively here too."""

    # Both documents were labelled by the frontend (their original filenames).
    label_1 = filename_1 or "Document 1"
    doc_1 = f"[Document: {label_1}]\n{pdf_text_1[:MAX_EXTRACTED_CHARS]}"

    if action == "compare" and pdf_text_2:
        label_2 = filename_2 or "Document 2"
        doc_2 = f"[Document: {label_2}]\n{pdf_text_2[:MAX_EXTRACTED_CHARS]}"
        return f"""You are an expert analyst. Carefully compare the two documents below and produce a
clear, well-organized comparison that highlights their similarities and differences.

Why are these documents different? Which is more useful for a student, and in what situation?

Document 1 ({label_1}):
{doc_1}

Document 2 ({label_2}):
{doc_2}

Comparison format:
1. Overview of each document (1-2 sentences each)
2. Similarities
3. Key differences (bulleted)
4. Conclusion — which is more suitable and for what purpose

Write in the language the documents are written in; if they are in different languages, write in the
language of Document 1."""

    if action == "ask":
        q = question.strip() or "Give a short overview of this document."
        return f"""You are a helpful study assistant. Answer the student's question about the document
below, staying grounded in the text. You may add a little helpful context, but don't contradict what
the document says.

{doc_1}

The student's question about this document:
{q}

Answer clearly and simply. Use bullet points, a short step-by-step breakdown, or an example where it
helps. Write in the same language as the question."""

    if action == "explain":
        return f"""You are a patient teacher. Explain the content of the document below in simple,
easy-to-understand language for a student who is new to the topic. Break down jargon and connect
the ideas clearly.

{doc_1}

Explain format:
1. What this is about, in plain words
2. The key ideas, explained simply
3. An easy example or analogy
4. Anything confusing, explained step by step

Write in the language of the document."""

    if action == "translate":
        lang = target_language.strip() or "English"
        return f"""Translate the document below into {lang}. Keep the meaning, tone, and structure
faithful to the original. Output only the translation, in that language's script.

{doc_1}"""

    # default: summarize
    return f"""You are an expert summarizer. Produce clear, well-organized study notes that summarize
the document below for a student preparing for exams.

{doc_1}

Summary format:
1. Short introduction
2. Main points (bulleted)
3. Key definitions / terms
4. Short conclusion

Write the ENTIRE summary in the language of the document."""


@app.route("/pdf-workspace-stream", methods=["POST"])
def pdf_workspace_stream():
    data = request.get_json(silent=True) or {}
    action = data.get("action", "").strip().lower()
    pdf_text_1 = data.get("pdf_text_1", "").strip()
    pdf_text_2 = data.get("pdf_text_2", "").strip()
    filename_1 = data.get("filename_1", "").strip()
    filename_2 = data.get("filename_2", "").strip()
    question = data.get("question", "").strip()
    target_language = data.get("target_language", "").strip()

    if action not in ("summarize", "ask", "explain", "translate", "compare"):
        return jsonify({"error": "Invalid action."}), 400

    if action == "compare":
        if not pdf_text_1 or not pdf_text_2:
            return jsonify({"error": "Compare ke liye do PDFs chahiye."}), 400
    elif not pdf_text_1:
        return jsonify({"error": "Pehle ek PDF upload karein."}), 400

    if action == "ask" and not question:
        return jsonify({"error": "Sawal khaali hai."}), 400

    prompt = _build_pdf_workspace_prompt(
        action, pdf_text_1, filename_1,
        pdf_text_2=pdf_text_2 or None,
        filename_2=filename_2 or None,
        question=question,
        target_language=target_language,
    )

    def generate():
        try:
            for chunk in ask_leo_stream(messages=[{"role": "user", "content": prompt}], stream=True):
                yield f"data: {json.dumps({'chunk': chunk})}\n\n"
            yield f"data: {json.dumps({'done': True})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return Response(stream_with_context(generate()), mimetype="text/event-stream")


# =========================
# AI BLOG WRITER (STREAMING)
# =========================

@app.route("/generate-blog-stream", methods=["POST"])
def generate_blog_stream():
    data = request.get_json(silent=True) or {}
    topic    = data.get("topic", "").strip()
    tone     = data.get("tone", "Professional").strip()
    length   = data.get("length", "Medium (600-800 words)").strip()
    keywords = data.get("keywords", "").strip()
    # Global default is English rather than Hindi (see note above).
    language = data.get("language", "English").strip()

    if not topic:
        return jsonify({"error": "Blog topic khaali hai."}), 400

    prompt = f"""
You are a professional content writer and SEO expert.

Write a blog post on this topic:

Topic: {topic}
Tone: {tone}
Length: {length}
Target keywords: {keywords if keywords else "N/A"}
Language: {language}

Write the ENTIRE blog in the language specified above — this could be English, Hindi, Spanish,
French, Arabic, Portuguese, Chinese, German, Japanese, Russian, Bengali, or any other language.
Follow it exactly, including using that language's own script.

Blog format:
1. Catchy SEO Title
2. Short Introduction (hook)
3. Main Body (with headings/subheadings, well structured)
4. Bullet points where useful
5. Conclusion
6. Call to Action

The blog should be engaging, SEO-friendly, and original.
"""

    def generate():
        try:
            for chunk in ask_leo_stream(messages=[{"role": "user", "content": prompt}], stream=True):
                yield f"data: {json.dumps({'chunk': chunk})}\n\n"
            yield f"data: {json.dumps({'done': True})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return Response(stream_with_context(generate()), mimetype="text/event-stream")

# =========================
# PDF DOWNLOAD
# =========================

# Unicode code-point ranges used to guess which script a piece of text is
# written in, so we can pick a font file that actually has glyphs for it.
# Add more (font_key, ttf_filename, range_check) entries here as you bundle
# more Noto Sans variants into static/fonts/.
_SCRIPT_FONT_RULES = [
    ("Devanagari",     "NotoSansDevanagari-Regular.ttf", lambda c: "\u0900" <= c <= "\u097F"),
    ("Bengali",        "NotoSansBengali-Regular.ttf",    lambda c: "\u0980" <= c <= "\u09FF"),
    ("Arabic",         "NotoSansArabic-Regular.ttf",     lambda c: "\u0600" <= c <= "\u06FF"),
    ("Hebrew",         "NotoSansHebrew-Regular.ttf",     lambda c: "\u0590" <= c <= "\u05FF"),
    ("Cyrillic",       "NotoSans-Regular.ttf",           lambda c: "\u0400" <= c <= "\u04FF"),
    ("CJK",            "NotoSansSC-Regular.ttf",         lambda c: "\u4E00" <= c <= "\u9FFF"),
    ("Japanese-Kana",  "NotoSansSC-Regular.ttf",         lambda c: "\u3040" <= c <= "\u30FF"),
    # Catch-all: any other non-ASCII character (accented Latin, Greek, etc.)
    # falls back to the broad-coverage NotoSans-Regular file.
    ("Latin-Extended",  "NotoSans-Regular.ttf",          lambda c: ord(c) > 127),
]

_registered_font_cache = {}


def pick_pdf_font(text):
    """Inspect `text` for non-ASCII scripts and return a registered
    ReportLab font name that can render it. Falls back to the built-in
    Helvetica font (Latin-only) if the text is plain ASCII, or if the
    matching Noto Sans font file isn't bundled under static/fonts/."""
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    FONT_DIR = os.path.join(BASE_DIR, "static", "fonts")

    sample = text[:3000]  # sampling is enough to detect the dominant script

    for key, filename, matcher in _SCRIPT_FONT_RULES:
        if not any(matcher(ch) for ch in sample):
            continue

        if key in _registered_font_cache:
            return _registered_font_cache[key]

        font_path = os.path.join(FONT_DIR, filename)
        if not os.path.exists(font_path):
            continue

        font_name = f"Font_{key}"
        try:
            pdfmetrics.registerFont(TTFont(font_name, font_path))
            _registered_font_cache[key] = font_name
            return font_name
        except Exception:
            continue

    return "Helvetica"


@app.route("/download-notes-pdf", methods=["POST"])
def download_notes_pdf():
    data = request.get_json(silent=True) or {}
    notes = data.get("notes", "").strip()

    if not notes:
        return jsonify({"error": "Notes khaali hai."}), 400

    # Unique filename per request so concurrent users never overwrite each other's PDF
    pdf_path = os.path.join(tempfile.gettempdir(), f"ai_notes_{uuid.uuid4().hex}.pdf")

    font_name = pick_pdf_font(notes)

    c = canvas.Canvas(pdf_path, pagesize=A4)
    width, height = A4
    x = 40
    y = height - 50

    c.setFont(font_name, 16)
    c.drawString(x, y, "AI Generated Notes")
    y -= 35

    c.setFont(font_name, 11)

    for paragraph in notes.split("\n"):
        lines = textwrap.wrap(paragraph, width=85)
        if not lines:
            y -= 15
            continue
        for line in lines:
            if y < 50:
                c.showPage()
                c.setFont(font_name, 11)
                y = height - 50
            c.drawString(x, y, line)
            y -= 17

    c.save()

    try:
        return send_file(
            pdf_path,
            as_attachment=True,
            download_name="ai_notes.pdf",
            mimetype="application/pdf"
        )
    finally:
        # Clean up the temp file after sending
        try:
            os.remove(pdf_path)
        except OSError:
            pass

@app.route("/download-notes-docx", methods=["POST"])
def download_notes_docx():
    data = request.get_json(silent=True) or {}
    notes = data.get("notes", "").strip()

    if not notes:
        return jsonify({"error": "Notes khaali hai."}), 400

    # Unique filename per request so concurrent users never overwrite each other's file
    docx_path = os.path.join(tempfile.gettempdir(), f"ai_notes_{uuid.uuid4().hex}.docx")

    # python-docx (already imported as DocxDocument for .docx upload reading)
    # fully supports WRITING too — a fresh Document() + add_paragraph() + save().
    doc = DocxDocument()

    # Bold title paragraph, then one paragraph per line (blank lines give spacing).
    title_run = doc.add_paragraph().add_run("AI Generated Notes")
    title_run.bold = True

    for paragraph in notes.split("\n"):
        doc.add_paragraph(paragraph)

    doc.save(docx_path)

    try:
        return send_file(
            docx_path,
            as_attachment=True,
            download_name="ai_notes.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    finally:
        # Clean up the temp file after sending
        try:
            os.remove(docx_path)
        except OSError:
            pass

# =========================
# DAILY STUDY STREAK
# =========================

def record_daily_activity(user):
    """Track the user's daily study streak (consecutive days of actual StudyAI
    activity — opening the dashboard, saving notes, or taking a quiz).

    Uses the UTC calendar date (datetime.utcnow().date()) to stay consistent
    with datetime.utcnow() everywhere else in the codebase:
      - already active today                 -> nothing changes
      - active yesterday (1-day gap)        -> current_streak += 1
      - first-ever activity or any gap      -> current_streak = 1
      - last_active_date is always set to today's UTC date whenever it changes
    """
    today = datetime.utcnow().date()

    if user.last_active_date == today:
        return  # already counted today — don't double-bump

    if user.last_active_date is not None and (today - user.last_active_date).days == 1:
        # (user.current_streak or 0) guards against NULL on pre-feature rows.
        user.current_streak = (user.current_streak or 0) + 1
    else:
        # First-ever activity (last_active_date is None) or a gap of 2+ days
        # (including clock-skew back) — either way a fresh streak starts today.
        user.current_streak = 1

    user.last_active_date = today
    db.session.commit()


# =========================
# ACCOUNT-SAVED NOTES + DASHBOARD
# (separate from the browser-local History panel in index.html — these
#  are cloud-saved and only available to logged-in users)
# =========================

@app.route("/api/notes", methods=["POST"])
@login_required
def api_save_note():
    data = request.get_json(silent=True) or {}
    content = (data.get("content") or "").strip()

    if not content:
        return jsonify({"error": "Notes khaali hai."}), 400

    record_daily_activity(current_user)  # saving notes counts as studying today

    note = Note(
        user_id=current_user.id,
        topic=(data.get("topic") or "Untitled")[:255],
        subject=(data.get("subject") or "")[:255],
        class_name=(data.get("class_name") or "")[:255],
        language=(data.get("language") or "English")[:64],
        content=content
    )
    db.session.add(note)
    db.session.commit()

    return jsonify({"id": note.id, "saved": True})


@app.route("/api/quiz-results", methods=["POST"])
@login_required
def api_save_quiz_result():
    data = request.get_json(silent=True) or {}

    try:
        score = int(data.get("score"))
        total = int(data.get("total"))
    except (TypeError, ValueError):
        return jsonify({"error": "Invalid score/total."}), 400

    if total <= 0 or score < 0 or score > total:
        return jsonify({"error": "Invalid score/total."}), 400

    record_daily_activity(current_user)  # finishing a quiz counts as studying today

    topic = (data.get("topic") or "Untitled")[:255]
    subject = (data.get("subject") or "")[:255]
    pct = round((score / total) * 100)

    result = QuizResult(
        user_id=current_user.id,
        topic=topic,
        subject=subject,
        score=score,
        total=total
    )
    db.session.add(result)

    # Update (or create) this topic's spaced-revision schedule based on how
    # the quiz went — see REVISION_STAGE_INTERVALS_DAYS for the rule.
    schedule = RevisionSchedule.query.filter_by(user_id=current_user.id, topic=topic).first()
    if not schedule:
        schedule = RevisionSchedule(user_id=current_user.id, topic=topic, subject=subject, interval_stage=0)
        db.session.add(schedule)

    max_stage = len(REVISION_STAGE_INTERVALS_DAYS) - 1
    if pct < 50:
        schedule.interval_stage = 0
    elif pct < 75:
        pass  # keep current stage — still shaky, review again at the same gap
    else:
        schedule.interval_stage = min(schedule.interval_stage + 1, max_stage)

    now = datetime.utcnow()
    schedule.subject = subject or schedule.subject
    schedule.last_score_pct = pct
    schedule.last_reviewed_at = now
    schedule.next_review_at = now + timedelta(days=REVISION_STAGE_INTERVALS_DAYS[schedule.interval_stage])

    db.session.commit()

    return jsonify({"id": result.id, "saved": True})


@app.route("/dashboard")
@login_required
def dashboard():
    record_daily_activity(current_user)  # opening the dashboard counts as studying today

    total_notes = Note.query.filter_by(user_id=current_user.id).count()
    recent_notes = (
        Note.query
        .filter_by(user_id=current_user.id)
        .order_by(Note.created_at.desc())
        .limit(20)
        .all()
    )

    quiz_results = (
        QuizResult.query
        .filter_by(user_id=current_user.id)
        .order_by(QuizResult.created_at.desc())
        .all()
    )

    total_quizzes = len(quiz_results)
    avg_score_pct = (
        round(sum(r.percent for r in quiz_results) / total_quizzes)
        if total_quizzes else None
    )

    # Weak-topic detection: average score per topic, weakest first. Only
    # topics with at least one attempt are considered; ties are broken by
    # attempt count (more data behind the average = more reliable signal).
    topic_stats = {}
    for r in quiz_results:
        key = r.topic or "Untitled"
        bucket = topic_stats.setdefault(key, {"scores": [], "count": 0})
        bucket["scores"].append(r.percent)
        bucket["count"] += 1

    weak_topics = sorted(
        (
            {
                "topic": topic,
                "avg": round(sum(b["scores"]) / len(b["scores"])),
                "attempts": b["count"]
            }
            for topic, b in topic_stats.items()
        ),
        key=lambda t: (t["avg"], -t["attempts"])
    )[:5]

    recent_quizzes = quiz_results[:10]

    # Spaced revision: topics due now (sorted most-overdue first), plus a
    # short look-ahead at what's coming up next.
    now = datetime.utcnow()
    due_revisions = (
        RevisionSchedule.query
        .filter_by(user_id=current_user.id)
        .filter(RevisionSchedule.next_review_at <= now)
        .order_by(RevisionSchedule.next_review_at.asc())
        .all()
    )
    upcoming_revisions = (
        RevisionSchedule.query
        .filter_by(user_id=current_user.id)
        .filter(RevisionSchedule.next_review_at > now)
        .order_by(RevisionSchedule.next_review_at.asc())
        .limit(5)
        .all()
    )

    return render_template(
        "dashboard.html",
        total_notes=total_notes,
        recent_notes=recent_notes,
        total_quizzes=total_quizzes,
        avg_score_pct=avg_score_pct,
        weak_topics=weak_topics,
        recent_quizzes=recent_quizzes,
        due_revisions=due_revisions,
        upcoming_revisions=upcoming_revisions,
        current_streak=current_user.current_streak,
        now=now
    )


@app.route("/dashboard/notes/<int:note_id>/delete", methods=["POST"])
@login_required
def dashboard_delete_note(note_id):
    note = Note.query.filter_by(id=note_id, user_id=current_user.id).first()
    if note:
        db.session.delete(note)
        db.session.commit()
    return redirect(url_for("dashboard"))


@app.route("/dashboard/revision/<int:schedule_id>/mark-reviewed", methods=["POST"])
@login_required
def dashboard_mark_reviewed(schedule_id):
    """Manual fallback for 'I reviewed this without taking a quiz' — advances
    the schedule the same way a strong quiz score would (assumes review went
    well; if it didn't, the next quiz attempt will reset the stage anyway)."""
    schedule = RevisionSchedule.query.filter_by(id=schedule_id, user_id=current_user.id).first()
    if schedule:
        max_stage = len(REVISION_STAGE_INTERVALS_DAYS) - 1
        schedule.interval_stage = min(schedule.interval_stage + 1, max_stage)
        now = datetime.utcnow()
        schedule.last_reviewed_at = now
        schedule.next_review_at = now + timedelta(days=REVISION_STAGE_INTERVALS_DAYS[schedule.interval_stage])
        db.session.commit()
    return redirect(url_for("dashboard"))


# =========================
# NOTE SHARING (public read-only links)
# =========================

@app.route("/dashboard/notes/<int:note_id>/toggle-share", methods=["POST"])
@login_required
def dashboard_toggle_share(note_id):
    """Flip the is_public flag for a note owned by the current user.
    If turning public ON and no share_token exists, generate one.
    Mirrors the ownership-check pattern of dashboard_delete_note."""
    note = Note.query.filter_by(id=note_id, user_id=current_user.id).first()
    if note:
        note.is_public = not note.is_public
        if note.is_public:
            note.ensure_share_token()  # generates token if missing
        db.session.commit()
    return redirect(url_for("dashboard"))


@app.route("/shared-notes/<share_token>")
def view_shared_note(share_token):
    """Public read-only view of a shared note.
    Requires BOTH share_token match AND is_public==True — if the owner
    later toggles the note back to private, the old token stops working.
    No @login_required — this must be accessible to logged-out visitors.
    Never exposes owning user's email or any other account details."""
    note = Note.query.filter_by(share_token=share_token, is_public=True).first()
    if not note:
        return render_template("shared_note.html", note=None), 404
    return render_template("shared_note.html", note=note)


# =========================
# NOTES SEARCH API
# =========================

@app.route("/api/notes/search")
@login_required
def api_notes_search():
    """
    Search the current user's saved notes by keyword across topic, subject, and content.
    Returns JSON with matching notes including a content snippet.
    """
    q = request.args.get("q", "").strip()
    
    # Handle empty/missing query - return empty results, not all notes
    if not q:
        return jsonify({"results": []})
    
    # Search across topic, subject, and content with case-insensitive matching
    # Only return notes belonging to the current logged-in user
    notes = (
        Note.query
        .filter_by(user_id=current_user.id)
        .filter(
            or_(
                Note.topic.ilike(f"%{q}%"),
                Note.subject.ilike(f"%{q}%"),
                Note.content.ilike(f"%{q}%")
            )
        )
        .order_by(Note.created_at.desc())
        .limit(30)
        .all()
    )
    
    results = []
    for note in notes:
        # Create a snippet - first ~200 chars of content
        snippet = note.content[:200]
        if len(note.content) > 200:
            snippet += "..."
        
        results.append({
            "id": note.id,
            "topic": note.topic,
            "subject": note.subject,
            "created_at": note.created_at.isoformat() if note.created_at else None,
            "snippet": snippet
        })
    
    return jsonify({"results": results})


# =========================
# USER DATA EXPORT / ACCOUNT DELETION
# =========================

@app.route("/api/export-data")
@login_required
def api_export_data():
    """
    Export all data for the current logged-in user as a downloadable JSON file.
    Includes profile (name, email, created_at), notes, quiz results, and revision schedules.
    Never exposes password_hash.
    """
    # Build the export data structure
    export_data = {
        "profile": {
            "id": current_user.id,
            "name": current_user.name,
            "email": current_user.email,
            "created_at": current_user.created_at.isoformat() if current_user.created_at else None,
        },
        "notes": [],
        "quiz_results": [],
        "revision_schedules": [],
        "exported_at": datetime.utcnow().isoformat(),
    }

    # Export notes
    notes = Note.query.filter_by(user_id=current_user.id).order_by(Note.created_at.desc()).all()
    for note in notes:
        export_data["notes"].append({
            "id": note.id,
            "topic": note.topic,
            "subject": note.subject,
            "class_name": note.class_name,
            "language": note.language,
            "content": note.content,
            "created_at": note.created_at.isoformat() if note.created_at else None,
        })

    # Export quiz results
    quiz_results = QuizResult.query.filter_by(user_id=current_user.id).order_by(QuizResult.created_at.desc()).all()
    for qr in quiz_results:
        export_data["quiz_results"].append({
            "id": qr.id,
            "topic": qr.topic,
            "subject": qr.subject,
            "score": qr.score,
            "total": qr.total,
            "percent": qr.percent,
            "created_at": qr.created_at.isoformat() if qr.created_at else None,
        })

    # Export revision schedules
    revision_schedules = RevisionSchedule.query.filter_by(user_id=current_user.id).order_by(RevisionSchedule.created_at.desc()).all()
    for rs in revision_schedules:
        export_data["revision_schedules"].append({
            "id": rs.id,
            "topic": rs.topic,
            "subject": rs.subject,
            "interval_stage": rs.interval_stage,
            "last_score_pct": rs.last_score_pct,
            "last_reviewed_at": rs.last_reviewed_at.isoformat() if rs.last_reviewed_at else None,
            "next_review_at": rs.next_review_at.isoformat() if rs.next_review_at else None,
            "created_at": rs.created_at.isoformat() if rs.created_at else None,
        })

    # Generate filename with user's email and timestamp
    safe_email = current_user.email.replace("@", "_at_").replace(".", "_")
    filename = f"studyai_export_{safe_email}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.json"

    # Return as downloadable JSON file
    json_str = json.dumps(export_data, indent=2, ensure_ascii=False)
    return Response(
        json_str,
        mimetype="application/json",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Content-Type": "application/json; charset=utf-8",
        }
    )


@app.route("/account/delete", methods=["POST"])
@login_required
def account_delete():
    """
    Permanently delete the current user's account and all associated data.
    Requires password confirmation in the request body.
    Deletes: Notes, QuizResults, RevisionSchedules, then the User itself.
    Logs out the user and redirects to home with a flash message.
    """
    # Get password from request body (JSON or form)
    data = request.get_json(silent=True) or {}
    password = data.get("password") or request.form.get("password")

    if not password:
        flash("Password confirmation is required.", "error")
        return redirect(url_for("home"))

    # Verify password
    if not current_user.check_password(password):
        flash("Incorrect password. Account not deleted.", "error")
        return redirect(url_for("home"))

    user_id = current_user.id

    # Delete all related data first (in correct order to avoid FK constraints)
    # Note: The models use lazy="dynamic" backrefs, so we can query and delete
    Note.query.filter_by(user_id=user_id).delete()
    QuizResult.query.filter_by(user_id=user_id).delete()
    RevisionSchedule.query.filter_by(user_id=user_id).delete()

    # Delete the user
    user = User.query.get(user_id)
    if user:
        db.session.delete(user)
        db.session.commit()

    # Log out the user
    logout_user()

    # Flash success message and redirect to home
    flash("Your account has been permanently deleted. All your data has been removed.", "success")
    return redirect(url_for("home"))


if __name__ == "__main__":
    # debug=False for anything beyond local development
    app.run(debug=False, port=5000)